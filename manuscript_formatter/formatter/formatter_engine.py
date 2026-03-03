from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional

from .nlp_section_detector import classify_document

# Built-in format templates
TEMPLATES = {
    "ieee": {
        "title":               {"size": 24, "bold": True,  "align": "center", "font": "Times New Roman", "space_after": 12},
        "authors":             {"size": 11, "bold": False, "align": "center", "italic": True, "font": "Times New Roman"},
        "abstract_heading":    {"size": 10, "bold": True,  "align": "left",   "font": "Times New Roman", "all_caps": True},
        "abstract_text":       {"size": 10, "italic": True,"align": "left",   "font": "Times New Roman"},
        "keywords_heading":    {"size": 10, "bold": True,  "align": "left",   "font": "Times New Roman"},
        "keywords_text":       {"size": 10, "italic": False,"align": "left",  "font": "Times New Roman"},
        "heading":             {"size": 12, "bold": True,  "align": "left",   "font": "Times New Roman", "all_caps": True},
        "subheading":          {"size": 11, "bold": True,  "italic": True, "align": "left", "font": "Times New Roman"},
        "body":                {"size": 10, "align": "justify", "font": "Times New Roman"},
        "references_heading":  {"size": 11, "bold": True,  "align": "left",   "font": "Times New Roman", "all_caps": True},
        "references_item":     {"size": 10, "align": "left", "font": "Times New Roman"},
        "figure_caption":      {"size": 9,  "align": "center", "italic": True, "font": "Times New Roman"},
        "table_caption":       {"size": 9,  "align": "center", "bold": True,  "font": "Times New Roman"},
        "acknowledgements":    {"size": 10, "bold": True,  "align": "left",   "font": "Times New Roman"},
        "appendix":            {"size": 11, "bold": True,  "align": "left",   "font": "Times New Roman"},
        "blank":               {},
    },
    "apa": {
        "title":               {"size": 14, "bold": True,  "align": "center", "font": "Times New Roman"},
        "authors":             {"size": 12, "bold": False, "align": "center", "font": "Times New Roman"},
        "abstract_heading":    {"size": 12, "bold": True,  "align": "center", "font": "Times New Roman"},
        "abstract_text":       {"size": 12, "align": "left", "font": "Times New Roman"},
        "keywords_heading":    {"size": 12, "italic": True,"align": "left",  "font": "Times New Roman"},
        "keywords_text":       {"size": 12, "align": "left", "font": "Times New Roman"},
        "heading":             {"size": 12, "bold": True,  "align": "center", "font": "Times New Roman"},
        "subheading":          {"size": 12, "bold": True,  "italic": True, "align": "left", "font": "Times New Roman"},
        "body":                {"size": 12, "align": "left", "font": "Times New Roman"},
        "references_heading":  {"size": 12, "bold": True,  "align": "center", "font": "Times New Roman"},
        "references_item":     {"size": 12, "align": "left", "font": "Times New Roman"},
        "figure_caption":      {"size": 12, "italic": True,"align": "left",  "font": "Times New Roman"},
        "table_caption":       {"size": 12, "align": "left", "font": "Times New Roman"},
        "acknowledgements":    {"size": 12, "bold": True,  "align": "center", "font": "Times New Roman"},
        "appendix":            {"size": 12, "bold": True,  "align": "center", "font": "Times New Roman"},
        "blank":               {},
    },
    "mla": {
        "title":               {"size": 12, "bold": False, "align": "center", "font": "Times New Roman"},
        "authors":             {"size": 12, "align": "left", "font": "Times New Roman"},
        "abstract_heading":    {"size": 12, "bold": True,  "align": "center", "font": "Times New Roman"},
        "abstract_text":       {"size": 12, "align": "left", "font": "Times New Roman"},
        "keywords_heading":    {"size": 12, "bold": True,  "align": "left",  "font": "Times New Roman"},
        "keywords_text":       {"size": 12, "align": "left", "font": "Times New Roman"},
        "heading":             {"size": 12, "bold": True,  "align": "left",  "font": "Times New Roman"},
        "subheading":          {"size": 12, "italic": True,"align": "left",  "font": "Times New Roman"},
        "body":                {"size": 12, "align": "left", "font": "Times New Roman"},
        "references_heading":  {"size": 12, "bold": True,  "align": "center","font": "Times New Roman"},
        "references_item":     {"size": 12, "align": "left", "font": "Times New Roman"},
        "figure_caption":      {"size": 12, "italic": True,"align": "left",  "font": "Times New Roman"},
        "table_caption":       {"size": 12, "align": "left", "font": "Times New Roman"},
        "acknowledgements":    {"size": 12, "align": "left", "font": "Times New Roman"},
        "appendix":            {"size": 12, "align": "left", "font": "Times New Roman"},
        "blank":               {},
    },
    "custom": {}   # populated dynamically from user input
}

DEFAULT_TEMPLATE = "ieee"


# ──────────────────────────────────────────────────────────────────────────────
# Apply formatting
# ──────────────────────────────────────────────────────────────────────────────

_ALIGN_MAP = {
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def apply_format(paragraph, style: dict):
    if not style:
        return

    paragraph.alignment = _ALIGN_MAP.get(
        style.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT
    )

    if style.get("space_after") is not None:
        paragraph.paragraph_format.space_after = Pt(style["space_after"])

    for run in paragraph.runs:
        f = run.font
        if style.get("size"):
            f.size = Pt(style["size"])
        if "bold" in style:
            f.bold = style["bold"]
        if "italic" in style:
            f.italic = style["italic"]

        # ✅ FIX: always explicitly set all_caps — never leave it inherited
        f.all_caps = style.get("all_caps", False)

        if style.get("font"):
            f.name = style["font"]
        if style.get("color"):
            r, g, b = style["color"]
            f.color.rgb = RGBColor(r, g, b)

    if not paragraph.runs and style.get("size"):
        run = paragraph.add_run()
        run.font.size = Pt(style["size"])
        run.font.all_caps = style.get("all_caps", False)  # ✅ FIX here too
        if style.get("font"):
            run.font.name = style["font"]
# ──────────────────────────────────────────────────────────────────────────────
# Main format_document function
# ──────────────────────────────────────────────────────────────────────────────

def format_document(
    input_path: str,
    output_path: str,
    template_name: str = DEFAULT_TEMPLATE,
    custom_template: Optional[dict] = None,
) -> dict:
    """
    Format a manuscript document using AI/NLP section detection.

    Parameters
    ----------
    input_path    : Path to uploaded .docx file.
    output_path   : Where to write the formatted .docx.
    template_name : One of "ieee", "apa", "mla", or "custom".
    custom_template: If template_name=="custom", provide the style dict here.

    Returns
    -------
    A report dict with section labels and confidence scores for every paragraph.
    """
    # ── load template ──
    if template_name == "custom" and custom_template:
        template = custom_template
    else:
        template = TEMPLATES.get(template_name, TEMPLATES[DEFAULT_TEMPLATE])

    # ── load document ──
    doc = Document(input_path)
    paragraphs = doc.paragraphs

    if not paragraphs:
        doc.save(output_path)
        return {"error": "Document is empty."}

    # ── NLP classification ──
    texts = [p.text for p in paragraphs]
    labels = classify_document(texts)

    # ── apply formatting ──
    report = []
    for para, label in zip(paragraphs, labels):
        style = template.get(label.label, template.get("body", {}))
        apply_format(para, style)
        report.append({
            "text_preview": para.text[:80],
            "detected_section": label.label,
            "confidence": label.confidence,
            "signals": label.signals,
        })

    doc.save(output_path)
    print(f"[formatter] Formatting complete → {output_path}")
    return {"template": template_name, "paragraphs": report}


# ──────────────────────────────────────────────────────────────────────────────
# CLI usage
# ──────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys, json
    if len(sys.argv) < 3:
        print("Usage: python formatter_engine.py input.docx output.docx [ieee|apa|mla]")
        sys.exit(1)
    inp, out = sys.argv[1], sys.argv[2]
    tmpl = sys.argv[3] if len(sys.argv) > 3 else DEFAULT_TEMPLATE
    result = format_document(inp, out, template_name=tmpl)
    print(json.dumps(result, indent=2))